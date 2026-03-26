from __future__ import annotations

from io import StringIO

from agent_app.tools.base import ToolResult
from agent_app.tools.filesystem import WorkspacePathMixin


class ReadPdfTool(WorkspacePathMixin):
    name = "read_pdf"
    description = '读取 PDF 文本内容，入参示例 {"path":"docs/report.pdf"}'
    input_schema = {"path": "string, .pdf file path under workspace"}
    example_input = '{"path":"docs/report.pdf"}'

    def run(self, tool_input: str) -> ToolResult:
        try:
            from pypdf import PdfReader

            payload = self._load_payload(tool_input)
            path = self._resolve_path(str(payload["path"]), must_exist=True)
            if path.suffix.lower() != ".pdf":
                raise ValueError("read_pdf 仅支持 .pdf 文件")

            reader = PdfReader(str(path))
            parts: list[str] = []
            for index, page in enumerate(reader.pages, start=1):
                text = (page.extract_text() or "").strip()
                if text:
                    parts.append(f"[Page {index}]\n{text}")

            if not parts:
                return self._ok("PDF 中未提取到可读文本。")
            return self._ok("\n\n".join(parts))
        except Exception as exc:
            return self._fail(exc)


class ReadXlsxTool(WorkspacePathMixin):
    name = "read_xlsx"
    description = (
        '读取 Excel 内容，入参示例 {"path":"data/demo.xlsx","sheet_name":"Sheet1"}'
    )
    input_schema = {
        "path": "string, .xlsx file path under workspace",
        "sheet_name": "string, optional worksheet name",
    }
    example_input = '{"path":"data/demo.xlsx","sheet_name":"Sheet1"}'

    def run(self, tool_input: str) -> ToolResult:
        workbook = None
        try:
            from openpyxl import load_workbook

            payload = self._load_payload(tool_input)
            path = self._resolve_path(str(payload["path"]), must_exist=True)
            if path.suffix.lower() != ".xlsx":
                raise ValueError("read_xlsx 仅支持 .xlsx 文件")

            sheet_name = str(payload["sheet_name"]) if "sheet_name" in payload else None
            workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
            worksheets = [workbook[sheet_name]] if sheet_name else list(workbook.worksheets)

            parts: list[str] = []
            for worksheet in worksheets:
                buffer = StringIO()
                buffer.write(f"[Sheet {worksheet.title}]\n")
                for row in worksheet.iter_rows(values_only=True):
                    values = ["" if cell is None else str(cell) for cell in row]
                    if any(values):
                        buffer.write("\t".join(values).rstrip())
                        buffer.write("\n")
                parts.append(buffer.getvalue().strip())

            content = "\n\n".join(part for part in parts if part)
            return self._ok(content or "Excel 中未读取到有效内容。")
        except Exception as exc:
            return self._fail(exc)
        finally:
            if workbook is not None:
                workbook.close()


class ReadWordTool(WorkspacePathMixin):
    name = "read_word"
    description = '读取 Word 内容，入参示例 {"path":"docs/notes.docx"}'
    input_schema = {"path": "string, .docx file path under workspace"}
    example_input = '{"path":"docs/notes.docx"}'

    def run(self, tool_input: str) -> ToolResult:
        try:
            from docx import Document

            payload = self._load_payload(tool_input)
            path = self._resolve_path(str(payload["path"]), must_exist=True)
            suffix = path.suffix.lower()
            if suffix == ".doc":
                raise ValueError("read_word 目前仅支持 .docx，不支持旧版 .doc")
            if suffix != ".docx":
                raise ValueError("read_word 仅支持 .docx 文件")

            document = Document(str(path))
            sections: list[str] = []

            paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
            paragraph_content = "\n".join(text for text in paragraphs if text)
            if paragraph_content:
                sections.append("[Paragraphs]\n" + paragraph_content)

            table_blocks: list[str] = []
            for index, table in enumerate(document.tables, start=1):
                rows: list[str] = []
                for row in table.rows:
                    values = [cell.text.strip() for cell in row.cells]
                    if any(values):
                        rows.append("\t".join(values))
                if rows:
                    table_blocks.append(f"[Table {index}]\n" + "\n".join(rows))
            if table_blocks:
                sections.append("\n\n".join(table_blocks))

            return self._ok("\n\n".join(sections) or "Word 中未读取到有效内容。")
        except Exception as exc:
            return self._fail(exc)
