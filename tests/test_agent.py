from __future__ import annotations

import json
import tempfile
import unittest
from pathlib import Path

from fastapi.testclient import TestClient

from agent_app.config import get_settings, resolve_model_config
from agent_app.core.agent import Agent
from agent_app.core.messages import Message
from agent_app.llm.base import LLMDecision, ToolSpec
from agent_app.llm.mock import MockLLM
from agent_app.main import (
    build_agent,
    build_memory_store,
    build_skill_registry,
    build_tool_registry,
)
from agent_app.memory.session_store import JsonSessionStore
from agent_app.memory.sqlite_store import SqliteMemoryStore
from agent_app.skills.loader import load_skill_from_markdown
from agent_app.skills.registry import SkillRegistry
from agent_app.tools.calculator import CalculatorTool
from agent_app.tools.document_readers import ReadPdfTool, ReadWordTool, ReadXlsxTool
from agent_app.tools.filesystem import (
    ListDirTool,
    MoveFileTool,
    ReadFileTool,
    RemoveFileTool,
    RenameFileTool,
    SearchInFilesTool,
    StatFileTool,
    WriteFileTool,
)
from agent_app.tools.registry import ToolRegistry
from agent_app.tools.time_tool import TimeNowTool


class InspectingLLM:
    def __init__(self) -> None:
        self.last_messages: list[Message] = []
        self.last_tools: list[ToolSpec] = []

    def decide(self, messages: list[Message], tools: list[ToolSpec]) -> LLMDecision:
        self.last_messages = list(messages)
        self.last_tools = list(tools)
        return LLMDecision(action="respond", content="ok")


class FixedDecisionLLM:
    def __init__(self, decision: LLMDecision) -> None:
        self._decision = decision
        self.last_tools: list[ToolSpec] = []

    def decide(self, messages: list[Message], tools: list[ToolSpec]) -> LLMDecision:
        self.last_tools = list(tools)
        return self._decision


class AgentFlowTestCase(unittest.TestCase):
    def test_agent_uses_calculator_tool(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            registry = ToolRegistry()
            registry.register(TimeNowTool())
            registry.register(CalculatorTool())

            agent = Agent(
                llm=MockLLM(),
                tool_registry=registry,
                session_store=JsonSessionStore(Path(tmpdir)),
            )

            reply = agent.run("2 * (3 + 4)", session_id="calc")

        self.assertIn("14", reply)

    def test_agent_reports_loaded_skills(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            registry = ToolRegistry()
            registry.register(TimeNowTool())
            skill_dir = Path(tmpdir) / "skills"
            skill_dir.mkdir(parents=True, exist_ok=True)
            (skill_dir / "inspector.md").write_text(
                (
                    "---\n"
                    "name: inspector\n"
                    "description: Inspect local content.\n"
                    "tools: read_file\n"
                    "---\n"
                    "Look at local files before answering.\n"
                ),
                encoding="utf-8",
            )

            agent = Agent(
                llm=MockLLM(),
                tool_registry=registry,
                session_store=JsonSessionStore(Path(tmpdir)),
                skill_registry=SkillRegistry.from_directory(skill_dir),
            )

            reply = agent.run(
                "当前有哪些技能",
                session_id="skills",
                skill_names=["inspector"],
            )

        self.assertIn("inspector", reply)

    def test_skill_allowed_tools_are_enforced(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            workspace = Path(tmpdir)
            registry = ToolRegistry()
            registry.register(ReadFileTool(workspace))
            registry.register(WriteFileTool(workspace))
            skill_dir = workspace / "skills"
            skill_dir.mkdir(parents=True, exist_ok=True)
            (skill_dir / "reader.md").write_text(
                (
                    "---\n"
                    "name: reader\n"
                    "description: Read only.\n"
                    "tools: read_file\n"
                    "---\n"
                    "Only use read_file.\n"
                ),
                encoding="utf-8",
            )

            llm = FixedDecisionLLM(
                LLMDecision(
                    action="tool",
                    tool_name="write_file",
                    tool_input=json.dumps({"path": "a.txt", "content": "blocked"}),
                )
            )
            agent = Agent(
                llm=llm,
                tool_registry=registry,
                session_store=JsonSessionStore(workspace / "sessions"),
                skill_registry=SkillRegistry.from_directory(skill_dir),
            )

            reply = agent.run("写点东西", session_id="restrict", skill_names=["reader"])

        self.assertIn("不在当前 skill 允许范围内", reply)
        self.assertEqual(["read_file"], [tool.name for tool in llm.last_tools])

    def test_relevant_memory_is_injected_into_runtime_messages(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            workspace = Path(tmpdir)
            registry = ToolRegistry()
            registry.register(TimeNowTool())
            memory_store = SqliteMemoryStore(workspace / "memory.db")
            memory_store.add_messages(
                "old-session",
                [Message(role="assistant", content="Alice likes structured notes.")],
            )
            llm = InspectingLLM()

            agent = Agent(
                llm=llm,
                tool_registry=registry,
                session_store=JsonSessionStore(workspace / "sessions"),
                memory_store=memory_store,
            )

            agent.run("请记住 Alice 的偏好", session_id="new-session")

        joined = "\n".join(
            message.content for message in llm.last_messages if message.role == "system"
        )
        self.assertIn("Relevant memory:", joined)
        self.assertIn("Alice likes structured notes.", joined)

    def test_tool_result_is_truncated(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            workspace = Path(tmpdir)
            registry = ToolRegistry()
            registry.register(ReadFileTool(workspace))
            long_text = "x" * 120
            (workspace / "long.txt").write_text(long_text, encoding="utf-8")

            llm = FixedDecisionLLM(
                LLMDecision(
                    action="tool",
                    tool_name="read_file",
                    tool_input=json.dumps({"path": "long.txt"}),
                )
            )
            agent = Agent(
                llm=llm,
                tool_registry=registry,
                session_store=JsonSessionStore(workspace / "sessions"),
                max_tool_iterations=1,
                max_tool_result_chars=20,
            )

            reply = agent.run("读取长文件", session_id="truncate")
            state = JsonSessionStore(workspace / "sessions").load("truncate")

        self.assertIn("已达到最大工具调用次数", reply)
        self.assertTrue(any("truncated" in message.content for message in state.messages))


class FileToolTestCase(unittest.TestCase):
    def test_file_tools_round_trip(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            workspace = Path(tmpdir)
            writer = WriteFileTool(workspace)
            reader = ReadFileTool(workspace)
            renamer = RenameFileTool(workspace)
            mover = MoveFileTool(workspace)
            remover = RemoveFileTool(workspace)

            write_result = writer.run(
                json.dumps({"path": "notes/todo.txt", "content": "hello agent"})
            )
            self.assertIn("写入文件", write_result.content)

            read_result = reader.run(json.dumps({"path": "notes/todo.txt"}))
            self.assertEqual("hello agent", read_result.content)

            rename_result = renamer.run(
                json.dumps({"path": "notes/todo.txt", "new_name": "done.txt"})
            )
            self.assertIn("重命名文件", rename_result.content)

            move_result = mover.run(
                json.dumps(
                    {
                        "source": "notes/done.txt",
                        "destination": "archive/done.txt",
                    }
                )
            )
            self.assertIn("移动文件", move_result.content)

            remove_result = remover.run(json.dumps({"path": "archive/done.txt"}))
            self.assertIn("删除文件", remove_result.content)
            self.assertFalse((workspace / "archive" / "done.txt").exists())

    def test_directory_listing_search_and_stat(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            workspace = Path(tmpdir)
            (workspace / "src").mkdir(parents=True, exist_ok=True)
            (workspace / "src" / "a.py").write_text(
                "print('hello')\n# TODO: refine\n",
                encoding="utf-8",
            )
            (workspace / "README.md").write_text("# Demo\n", encoding="utf-8")

            list_result = ListDirTool(workspace).run(
                json.dumps({"path": ".", "recursive": True})
            )
            search_result = SearchInFilesTool(workspace).run(
                json.dumps({"path": ".", "pattern": "TODO", "glob": "*.py"})
            )
            stat_result = StatFileTool(workspace).run(json.dumps({"path": "README.md"}))

        self.assertIn("[file] README.md", list_result.content)
        self.assertIn("src/a.py:2", search_result.content)
        stat_payload = json.loads(stat_result.content)
        self.assertEqual("README.md", stat_payload["path"])
        self.assertFalse(stat_payload["is_dir"])

    def test_file_tools_reject_paths_outside_workspace(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            reader = ReadFileTool(Path(tmpdir))
            result = reader.run(json.dumps({"path": "../outside.txt"}))
            self.assertIn("工作目录内", result.content)

    def test_tool_specs_include_schema_and_examples(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            settings = get_settings(model_provider="mock")
            settings.workspace_root = Path(tmpdir)
            registry = build_tool_registry(settings)
            tools = {tool.name: tool for tool in registry.specs()}

        self.assertIn("path", tools["read_file"].input_schema)
        self.assertTrue(tools["search_in_files"].example_input)


class DocumentReaderToolTestCase(unittest.TestCase):
    def test_read_pdf_tool_extracts_text(self) -> None:
        from reportlab.pdfgen import canvas

        with tempfile.TemporaryDirectory() as tmpdir:
            workspace = Path(tmpdir)
            pdf_path = workspace / "docs" / "sample.pdf"
            pdf_path.parent.mkdir(parents=True, exist_ok=True)

            pdf = canvas.Canvas(str(pdf_path))
            pdf.drawString(72, 720, "Hello PDF")
            pdf.drawString(72, 700, "Second line")
            pdf.save()

            reader = ReadPdfTool(workspace)
            result = reader.run(json.dumps({"path": "docs/sample.pdf"}))

        self.assertIn("Hello PDF", result.content)
        self.assertIn("Second line", result.content)

    def test_read_xlsx_tool_extracts_cells(self) -> None:
        from openpyxl import Workbook

        with tempfile.TemporaryDirectory() as tmpdir:
            workspace = Path(tmpdir)
            xlsx_path = workspace / "data" / "sample.xlsx"
            xlsx_path.parent.mkdir(parents=True, exist_ok=True)

            workbook = Workbook()
            sheet = workbook.active
            sheet.title = "Summary"
            sheet.append(["name", "score"])
            sheet.append(["alice", 95])
            workbook.save(xlsx_path)
            workbook.close()

            reader = ReadXlsxTool(workspace)
            result = reader.run(json.dumps({"path": "data/sample.xlsx"}))

        self.assertIn("[Sheet Summary]", result.content)
        self.assertIn("alice\t95", result.content)

    def test_read_word_tool_extracts_paragraphs_and_tables(self) -> None:
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            workspace = Path(tmpdir)
            docx_path = workspace / "docs" / "sample.docx"
            docx_path.parent.mkdir(parents=True, exist_ok=True)

            document = Document()
            document.add_paragraph("Hello Word")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "name"
            table.cell(0, 1).text = "role"
            table.cell(1, 0).text = "alice"
            table.cell(1, 1).text = "agent"
            document.save(docx_path)

            reader = ReadWordTool(workspace)
            result = reader.run(json.dumps({"path": "docs/sample.docx"}))

        self.assertIn("Hello Word", result.content)
        self.assertIn("alice\tagent", result.content)


class SkillAndMemoryTestCase(unittest.TestCase):
    def test_load_skill_from_markdown(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            skill_path = Path(tmpdir) / "reader.md"
            skill_path.write_text(
                (
                    "---\n"
                    "name: reader\n"
                    "description: Read local documents.\n"
                    "tools: read_file, read_pdf\n"
                    "---\n"
                    "Use file tools before answering.\n"
                ),
                encoding="utf-8",
            )

            skill = load_skill_from_markdown(skill_path)

        self.assertEqual("reader", skill.name)
        self.assertEqual(("read_file", "read_pdf"), skill.allowed_tools)
        self.assertIn("Use file tools", skill.prompt)

    def test_build_skill_registry_and_memory_store_from_settings(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            skill_dir = root / "skills"
            skill_dir.mkdir(parents=True, exist_ok=True)
            (skill_dir / "writer.md").write_text(
                (
                    "---\n"
                    "name: writer\n"
                    "description: Write files safely.\n"
                    "---\n"
                    "Prefer write_file for edits.\n"
                ),
                encoding="utf-8",
            )
            settings = get_settings(model_provider="mock")
            settings.skills_dir = skill_dir
            settings.memory_db_path = root / "memory.db"

            registry = build_skill_registry(settings)
            memory_store = build_memory_store(settings)
            memory_store.add_messages(
                "session-a",
                [Message(role="assistant", content="Remember the release checklist.")],
            )
            hits = memory_store.search("release")

            self.assertEqual(["writer"], [skill.name for skill in registry.list_skills()])
            self.assertEqual(1, len(hits))
            self.assertEqual(["session-a"], memory_store.list_sessions())
            self.assertGreater(memory_store.delete_session("session-a"), 0)


class FastAPITestCase(unittest.TestCase):
    def test_fastapi_endpoints_work(self) -> None:
        from agent_app.api.server import create_app

        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            settings = get_settings(model_provider="mock")
            settings.workspace_root = root
            settings.session_store_dir = root / "sessions"
            settings.skills_dir = root / "skills"
            settings.memory_db_path = root / "memory.db"
            settings.skills_dir.mkdir(parents=True, exist_ok=True)
            (settings.skills_dir / "helper.md").write_text(
                (
                    "---\n"
                    "name: helper\n"
                    "description: Helpful skill.\n"
                    "tools: calculator\n"
                    "---\n"
                    "Use calculator when useful.\n"
                ),
                encoding="utf-8",
            )

            client = TestClient(create_app(settings))
            health = client.get("/health")
            config = client.get("/config")
            tools = client.get("/tools")
            skills = client.get("/skills")
            chat = client.post(
                "/chat",
                json={"message": "2 * (3 + 4)", "session_id": "api", "skills": ["helper"]},
            )
            memories = client.get("/memories/search", params={"q": "14"})
            sessions = client.get("/sessions")
            session_detail = client.get("/sessions/api")
            memory_delete = client.delete("/memories/session/api")
            session_delete = client.delete("/sessions/api")

        self.assertEqual(200, health.status_code)
        self.assertEqual({"status": "ok"}, health.json())
        self.assertEqual(200, config.status_code)
        self.assertEqual("mock", config.json()["model_provider"])
        self.assertEqual(200, tools.status_code)
        self.assertTrue(any(tool["name"] == "list_dir" for tool in tools.json()))
        self.assertTrue(any(tool["input_schema"] for tool in tools.json()))
        self.assertEqual(200, skills.status_code)
        self.assertEqual("helper", skills.json()[0]["name"])
        self.assertEqual(200, chat.status_code)
        self.assertIn("14", chat.json()["response"])
        self.assertEqual(200, memories.status_code)
        self.assertTrue(len(memories.json()) >= 1)
        self.assertEqual(200, sessions.status_code)
        self.assertIn("api", sessions.json()["sessions"])
        self.assertEqual(200, session_detail.status_code)
        self.assertEqual("api", session_detail.json()["session_id"])
        self.assertEqual(200, memory_delete.status_code)
        self.assertEqual(200, session_delete.status_code)


class ModelConfigTestCase(unittest.TestCase):
    def test_openai_preset(self) -> None:
        config = resolve_model_config(provider="openai", api_key_override="test-openai")
        self.assertEqual("openai", config.provider)
        self.assertEqual("gpt-4.1-mini", config.model_name)
        self.assertEqual("https://api.openai.com/v1", config.base_url)
        config.validate()

    def test_gemini_preset(self) -> None:
        config = resolve_model_config(provider="gemini", api_key_override="test-gemini")
        self.assertEqual("gemini", config.provider)
        self.assertEqual("gemini-2.5-flash", config.model_name)
        self.assertEqual(
            "https://generativelanguage.googleapis.com/v1beta/openai/",
            config.base_url,
        )
        config.validate()

    def test_glm_alias_and_preset(self) -> None:
        config = resolve_model_config(provider="glm-4.7", api_key_override="test-glm")
        self.assertEqual("glm", config.provider)
        self.assertEqual("glm-4.7", config.model_name)
        self.assertEqual("https://api.z.ai/api/paas/v4/", config.base_url)
        config.validate()

    def test_settings_builds_registry_with_tools(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            settings = get_settings(model_provider="mock")
            settings.workspace_root = Path(tmpdir)
            registry = build_tool_registry(settings)
            tool_names = {spec.name for spec in registry.specs()}

        self.assertTrue(
            {
                "read_file",
                "write_file",
                "rename_file",
                "move_file",
                "rm_file",
                "list_dir",
                "search_in_files",
                "stat_file",
                "read_pdf",
                "read_xlsx",
                "read_word",
            }.issubset(tool_names)
        )

    def test_build_agent_with_sqlite_memory(self) -> None:
        with tempfile.TemporaryDirectory() as tmpdir:
            root = Path(tmpdir)
            settings = get_settings(model_provider="mock")
            settings.workspace_root = root
            settings.session_store_dir = root / "sessions"
            settings.skills_dir = root / "skills"
            settings.memory_db_path = root / "memory.db"
            agent = build_agent(settings)
            reply = agent.run("现在几点", session_id="time")

        self.assertTrue(reply)


if __name__ == "__main__":
    unittest.main()
